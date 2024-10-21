#!/usr/bin/env python
# coding: utf-8

import os
from os.path import join
import docx
#from markupsafe import escape # which should be used?
from html import escape
from spire.doc import * # import Document() class
from spire.doc.common import * # import Document() class

# PIPELINE FUNCTIONS

## Docx reader
def read_docx_to_string(docx_file):
    # Open the .docx file
    doc = docx.Document(docx_file)
    
    # Extract text from each paragraph and join into a single string
    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    
    return full_text


## Highlight sentences
def highlight_sentences(sentences, highlight_label, highlight_color="yellow"):
    # Escape HTML characters in sentences and highlight those containing the specified label
    highlighted_sentences = [
        f'<span style="background-color:{highlight_color};">{escape(sentence.get("sentence"))}</span>'
        if sentence.get("label") == highlight_label else escape(sentence.get("sentence"))
        for sentence in sentences
    ]
    
    # Join the sentences back together into a single string
    highlighted_text = "".join(highlighted_sentences)
    
    return highlighted_text


## Convert to HTML
def generate_html(highlighted_text):
    
    # Wrap the highlighted text in basic HTML structure
    html_content = f"""<!DOCTYPE html>
    <head>
    <meta charset="UTF-8">
    </head>
    <body>
        <pre>{highlighted_text}</pre>
    </body>
    </html>"""
    
    return html_content


# Function to convert HTML file to DOCX
def convert_html_to_docx(html_filepath):

    output_file_path = html_filepath.replace(".html", ".docx")
            
    # Create an object of the Document class
    document = Document()
    
    try:
        # Load the HTML file
        document.LoadFromFile(html_filepath, FileFormat.Html, XHTMLValidationType.none)
        
        # Save the HTML file to a .docx file
        document.SaveToFile(output_file_path, FileFormat.Docx2019)
        
        print(f"Successfully converted {html_filepath} to DOCX")
    except Exception as e:
        print(f"Failed to convert {html_filepath}: {e}")
    finally:
        # Close the document
        document.Close()